import requests
import os
import json
from collections import defaultdict
from dotenv import load_dotenv
from docx import Document
import networkx as nx
import matplotlib.pyplot as plt
import numpy as np
import matplotlib.colors as mcolors
from matplotlib.patches import Circle
from matplotlib.collections import PatchCollection
from textblob import TextBlob
from rake_nltk import Rake
import nltk

# Load environment variables from .env file
load_dotenv()

class BrainstormAgent:
    @staticmethod
    def get_user_input():
        """Get user input for the brainstorming topic, output option, and language."""
        language = input("Choose language (1: English, 2: German): ")
        topic = input("Enter your brainstorming prompt: " if language == "1" else "Geben Sie Ihr Brainstorming-Thema ein: ")
        output_option = input("Choose output option (1: Generate mind map, 2: Print ideas to Word file, 3: Export as JSON): " if language == "1" else "Wählen Sie die Ausgabeoption (1: Mindmap generieren, 2: Ideen in Word-Datei drucken, 3: Als JSON exportieren): ")
        return topic, output_option, language

    def __init__(self, api_key):
        self.ideas = []
        self.mind_map = defaultdict(list)
        self.api_key = api_key
        self.base_url = "https://openrouter.ai/api/v1/chat/completions"
        
        # Download NLTK data if not already present
        try:
            nltk.data.find('corpora/stopwords')
        except LookupError:
            print("Downloading NLTK stopwords...")
            nltk.download('stopwords', quiet=True)
        
        self.rake = Rake()

    def analyze_sentiment(self, text):
        """Analyze the sentiment of the given text."""
        blob = TextBlob(text)
        return blob.sentiment.polarity

    def extract_keywords(self, text):
        """Extract keywords from the given text."""
        self.rake.extract_keywords_from_text(text)
        return self.rake.get_ranked_phrases()[:5]  # Return top 5 keywords

    def generate_ideas(self, topic, num_ideas=10, language="1"):
        """Generate a list of ideas related to the given topic using OpenRouter API."""
        if language == "1":
            prompt = f"Generate {num_ideas} innovative ideas related to {topic}. Provide each idea as a short phrase or sentence."
        else:
            prompt = f"Generiere {num_ideas} innovative Ideen zu {topic}. Gib jede Idee als kurze Phrase oder einen Satz an."
        
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "anthropic/claude-3.5-sonnet",
            "messages": [{"role": "user", "content": prompt}]
        }
        
        response = requests.post(self.base_url, json=data, headers=headers)
        response.raise_for_status()
        
        result = response.json()
        ideas_text = result['choices'][0]['message']['content']
        
        raw_ideas = [idea.strip() for idea in ideas_text.split('\n') if idea.strip()]
        self.ideas = []
        for idea in raw_ideas[:num_ideas]:
            sentiment = self.analyze_sentiment(idea)
            keywords = self.extract_keywords(idea)
            self.ideas.append({
                "text": idea,
                "sentiment": sentiment,
                "keywords": keywords
            })
        return self.ideas

    def create_mind_map(self):
        """Create a mind map from the generated ideas."""
        for idea in self.ideas:
            words = idea.split()
            for word in words:
                self.mind_map[word].append(idea)
        return dict(self.mind_map)

    def print_mind_map(self):
        """Print the mind map in a simple text format."""
        for key, values in self.mind_map.items():
            print(f"{key}:")
            for value in values:
                print(f"  - {value}")
            print()

    def create_networkx_mind_map(self, output_file="mind_map.png"):
        """Create a modern and visually appealing NetworkX mind map from the generated ideas."""
        G = nx.Graph()
        
        # Color palette
        colors = plt.cm.viridis(np.linspace(0, 1, 10))
        
        # Add nodes and edges
        main_topic = list(self.mind_map.keys())[0]
        G.add_node(main_topic, color=colors[0], size=8000)
        
        for i, (subtopic, ideas) in enumerate(self.mind_map[main_topic].items()):
            subtopic_color = colors[(i + 1) % len(colors)]
            G.add_node(subtopic, color=subtopic_color, size=6000)
            G.add_edge(main_topic, subtopic)
            
            for j, idea in enumerate(ideas):
                idea_color = mcolors.to_rgba(subtopic_color, alpha=0.7)
                G.add_node(idea, color=idea_color, size=4000)
                G.add_edge(subtopic, idea)
        
        # Set up the plot
        plt.figure(figsize=(24, 18), facecolor='#f0f0f0')
        pos = nx.spring_layout(G, k=0.5, iterations=50)
        
        # Draw edges
        nx.draw_networkx_edges(G, pos, edge_color='#d0d0d0', width=1.5, alpha=0.7)
        
        # Draw nodes with gradients
        node_colors = [G.nodes[node]['color'] for node in G.nodes()]
        node_sizes = [G.nodes[node]['size'] for node in G.nodes()]
        
        for node, (x, y) in pos.items():
            color = G.nodes[node]['color']
            size = G.nodes[node]['size'] / 1000
            circle = Circle((x, y), size/200, facecolor="none")
            plt.gca().add_patch(circle)
            plt.gca().add_collection(PatchCollection([circle], facecolors=[color]))
        
        # Draw labels
        labels = {node: self.wrap_text(node, 20) for node in G.nodes()}
        font_sizes = {node: 16 if G.nodes[node]['size'] > 6000 else 12 if G.nodes[node]['size'] > 4000 else 8 for node in G.nodes()}
        for node, (x, y) in pos.items():
            plt.text(x, y, labels[node], fontsize=font_sizes[node], ha='center', va='center', wrap=True, fontweight='bold', color='#303030')
        
        # Add a subtle grid
        plt.grid(color='#e0e0e0', linestyle='--', linewidth=0.5, alpha=0.5)
        
        # Save the plot
        plt.axis('off')
        plt.tight_layout()
        plt.savefig(output_file, format="png", dpi=300, bbox_inches="tight", facecolor='#f0f0f0', edgecolor='none')
        plt.close()
        
        print(f"Modern mind map has been saved as {output_file}")

    @staticmethod
    def wrap_text(text, max_width):
        """Wrap text to a maximum width."""
        words = text.split()
        lines = []
        current_line = []
        current_length = 0

        for word in words:
            if current_length + len(word) <= max_width:
                current_line.append(word)
                current_length += len(word) + 1
            else:
                lines.append(' '.join(current_line))
                current_line = [word]
                current_length = len(word)

        if current_line:
            lines.append(' '.join(current_line))

        return '\n'.join(lines)


    def write_ideas_to_word(self, filename="brainstorm_ideas.docx"):
        """Write generated ideas to a Word document."""
        doc = Document()
        doc.add_heading('Brainstorming Ideas', 0)
        for idea in self.ideas:
            doc.add_paragraph(idea, style='List Bullet')
        doc.save(filename)
        print(f"Ideas have been written to {filename}")

    def generate_mind_map(self, topic, language="1"):
        """Generate a comprehensive mind map structure using OpenRouter API."""
        if language == "1":
            prompt = f"""Create a detailed mind map for the topic: {topic}.
            The mind map should have the following structure:
            1. Main topic
            2. At least 5 subtopics
            3. For each subtopic, provide at least 3 related ideas or concepts
            
            Provide the output as a JSON object where:
            - The key is the main topic
            - The value is a dictionary of subtopics
            - Each subtopic has an array of related ideas

            Example format:
            {{
                "Main Topic": {{
                    "Subtopic 1": ["Idea 1", "Idea 2", "Idea 3"],
                    "Subtopic 2": ["Idea 1", "Idea 2", "Idea 3"],
                    ...
                }}
            }}
            """
        else:
            prompt = f"""Erstelle eine detaillierte Mindmap für das Thema: {topic}.
            Die Mindmap sollte folgende Struktur haben:
            1. Hauptthema
            2. Mindestens 5 Unterthemen
            3. Für jedes Unterthema mindestens 3 verwandte Ideen oder Konzepte
            
            Gib die Ausgabe als JSON-Objekt an, wobei:
            - Der Schlüssel das Hauptthema ist
            - Der Wert ein Wörterbuch der Unterthemen ist
            - Jedes Unterthema ein Array verwandter Ideen hat

            Beispielformat:
            {{
                "Hauptthema": {{
                    "Unterthema 1": ["Idee 1", "Idee 2", "Idee 3"],
                    "Unterthema 2": ["Idee 1", "Idee 2", "Idee 3"],
                    ...
                }}
            }}
            """
        
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "anthropic/claude-3.5-sonnet",
            "messages": [{"role": "user", "content": prompt}]
        }
        
        try:
            response = requests.post(self.base_url, json=data, headers=headers)
            response.raise_for_status()
            
            result = response.json()
            mind_map_text = result['choices'][0]['message']['content']
            
            # Extract the JSON part from the response
            import json
            start = mind_map_text.find('{')
            end = mind_map_text.rfind('}') + 1
            mind_map_json = mind_map_text[start:end]
            
            self.mind_map = json.loads(mind_map_json)
            return self.mind_map
        except requests.exceptions.RequestException as e:
            print(f"An error occurred: {e}")
            if hasattr(e, 'response') and e.response is not None:
                print(f"Response status code: {e.response.status_code}")
                print(f"Response content: {e.response.text}")
            return None

# Example usage
if __name__ == "__main__":
    api_key = os.getenv("OPENROUTER_API_KEY")
    if not api_key:
        raise ValueError("Please set the OPENROUTER_API_KEY in the .env file")
    
    agent = BrainstormAgent(api_key)
    topic, output_option, language = BrainstormAgent.get_user_input()
    
    if output_option == "1":
        print(f"\nGenerating mind map for: {topic}" if language == "1" else f"\nMindmap wird generiert für: {topic}")
        mind_map = agent.generate_mind_map(topic, language)
        if mind_map:
            print("\nMind Map structure:" if language == "1" else "\nMindmap-Struktur:")
            agent.print_mind_map()
            agent.create_networkx_mind_map()
        else:
            print("Failed to generate mind map. Please check your API key and try again." if language == "1" else "Fehler beim Generieren der Mindmap. Bitte überprüfen Sie Ihren API-Schlüssel und versuchen Sie es erneut.")
    elif output_option == "2":
        print(f"\nGenerating ideas for: {topic}" if language == "1" else f"\nIdeen werden generiert für: {topic}")
        ideas = agent.generate_ideas(topic, language=language)
        print("\nGenerated ideas:" if language == "1" else "\nGenerierte Ideen:")
        for idea in ideas:
            print(f"- {idea['text']}")
            print(f"  Sentiment: {idea['sentiment']:.2f}")
            print(f"  Keywords: {', '.join(idea['keywords'])}")
        agent.write_ideas_to_word()
    elif output_option == "3":
        print(f"\nGenerating mind map for: {topic}" if language == "1" else f"\nMindmap wird generiert für: {topic}")
        mind_map = agent.generate_mind_map(topic, language)
        if mind_map:
            agent.export_mind_map_json()
        else:
            print("Failed to generate mind map. Please check your API key and try again." if language == "1" else "Fehler beim Generieren der Mindmap. Bitte überprüfen Sie Ihren API-Schlüssel und versuchen Sie es erneut.")
    else:
        print("Invalid option selected. Please run the script again and choose either 1, 2, or 3." if language == "1" else "Ungültige Option ausgewählt. Bitte führen Sie das Skript erneut aus und wählen Sie entweder 1, 2 oder 3.")
    def export_mind_map_json(self, filename="mind_map.json"):
        """Export the mind map as a JSON file."""
        with open(filename, 'w') as f:
            json.dump(self.mind_map, f, indent=2)
        print(f"Mind map has been exported as {filename}")
# Rename this file to ideaforge_ai.py
